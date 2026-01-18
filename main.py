# main.py – BN News DOCX generátor (FastAPI)
# ------------------------------------------
from fastapi import FastAPI, Response, HTTPException
from pydantic import BaseModel
from urllib.parse import quote, urlparse
from readability import Document as ReadabilityDoc
from lxml import html
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_BREAK
import os, io, re, datetime, requests, pandas as pd
import trafilatura
import google.generativeai as genai
import json

# ===== Konfiguráció =====
TEMPLATE_PATH = "ceges_sablon.docx"
REQUIRED_COLS = {"Rovat", "Link"}
APP_SECRET = "007"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

app = FastAPI()

# ===== Payload osztályok (A hiba elkerülése végett elöl) =====
class Payload(BaseModel):
    sheet_id: str
    worksheet: str
    rovat: str
    secret: str | None = None

class ChatPayload(BaseModel):
    sheet_id: str
    worksheet: str
    rovat: str
    query: str | None = ""
    secret: str | None = None

# ===== Segédek =====
def csv_url(sheet_id: str, sheet_name: str) -> str:
    return (
        f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?"
        f"tqx=out:csv&sheet={quote(sheet_name)}"
    )

def norm_space(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()

# mondatzáró jel
SENT_END_RE = re.compile(r'[.!?…]"?$')

# reklám/junk minták - KIEGÉSZÍTVE AZ ECONOMX SALLANGOKKAL
AD_PATTERNS = [
    r"^\s*hirdet[ée]s\b",
    r"^\s*szponzor[áa]lt\b",
    r"^\s*t[áa]mogatott tartalom\b",
    r"^\s*aj[áa]nl[oó]\b",
    r"^\s*kapcsol[óo]d[óo].*",
    r"^\s*olvasta m[áa]r\??",
    r"^\s*promo",
    r"^\s*advertisement\b",
    r"^\s*sponsored\b",
    r"^source:\s*.+$",
    r".*\bc[íi]mlapk[ée]p\b.*",
    r".*\bbor[íi]t[óo]k[ée]p\b.*",
    r".*\b(getty images|shutterstock|reuters|associated press|ap photo|afp|epa)\b.*",
    r"^\s*back to intro\b",
    r"^\s*read article\b",
    r"^érdekesnek találta.*hírlevelünkre",
    r"^\s*hírlev[ée]l",
    r"^\s*kapcsol[óo]d[óo] cikk(ek)?\b",
    r"^\s*fot[óo]gal[ée]ria\b",
    r"^\s*tov[áa]bbi (h[íi]reink|cikkek)\b",
    r"^\s*Csapjunk bele a közepébe",
    r"A cikk elkészítésében .* Alrite .* alkalmazás támogatta a munkánkat\.?$",
    r".* – Fotó: .+$",
    # Economx specifikus
    r"A gazdaság és az üzleti élet legfrissebb hírei az Economx.hu hírlevelében",
    r"Küldtünk Önnek egy emailt!",
    r"feliratkozása megerősítéséhez"
]
JUNK_RE = re.compile("|".join(AD_PATTERNS), flags=re.IGNORECASE)

def is_sentence_like(s: str) -> bool:
    s = s.strip()
    # A felsorolásokat (• vagy -) is mondatnak tekintjük, hogy ne vesszenek el
    if s.startswith("•") or s.startswith("- "):
        return True
    return bool(SENT_END_RE.search(s)) or len(s) > 200

def clean_and_merge(paras: list[str]) -> list[str]:
    """Bekezdések tisztítása és teljes mondatokra fűzése."""
    lines = []
    for p in paras:
        t = norm_space(p)
        if not t:
            continue
        if JUNK_RE.search(t):
            continue
        # Felsorolásokat akkor is megtartjuk, ha rövidek
        if t.startswith("•") or t.startswith("- "):
            lines.append(t)
            continue
        # nagyon rövid, feltehetően alcím – kihagyjuk
        if len(t) < 35 and not t.endswith(":"):
            continue
        lines.append(t)

    merged, buf = [], ""
    for t in lines:
        # Ha felsorolás jön, a puffert lezárjuk és a listatagot külön tesszük
        if t.startswith("•") or t.startswith("- "):
            if buf:
                merged.append(buf)
                buf = ""
            merged.append(t)
            continue

        buf = f"{buf} {t}".strip() if buf else t
        if is_sentence_like(buf):
            merged.append(buf)
            buf = ""
    # ha maradt valami hosszabb a pufferben, engedjük át
    if buf and len(buf) > 60:
        merged.append(buf)
    merged = [m for m in merged if not JUNK_RE.search(m)]
    return merged

# (add_bm, add_link, hu_date függvények változatlanok...)
def add_bm(paragraph, name: str):
    run = paragraph.add_run()
    r = run._r
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), "1")
    bs.set(qn("w:name"), name)
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), "1")
    r.append(bs)
    r.append(be)

def add_link(paragraph, text: str, anchor: str):
    h = OxmlElement("w:hyperlink")
    h.set(qn("w:anchor"), anchor)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    rPr.append(c)
    t = OxmlElement("w:t")
    t.text = text
    r.append(rPr)
    r.append(t)
    h.append(r)
    paragraph._p.append(h)

def hu_date(d: datetime.date) -> str:
    return d.strftime("%Y.%m.%d.")

# ===== Cikk kinyerés =====
def read_paras(url: str):
    # 1) Readability
    try:
        r = requests.get(url, timeout=25, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
        rd = ReadabilityDoc(r.text)
        title = (rd.short_title() or "").strip()
        root = html.fromstring(rd.summary())
        # KIEGÉSZÍTVE: li (listaelemek) kinyerése és jelölése
        paras = []
        for el in root.xpath(".//p | .//li"):
            text = norm_space(el.text_content())
            if el.tag == "li":
                text = f"• {text}"
            paras.append(text)
        
        paras = [p for p in paras if p and not JUNK_RE.search(p)]
        cleaned = clean_and_merge(paras)
        if cleaned:
            return title, cleaned
    except Exception:
        pass
    # 2) trafilatura
    try:
        dl = trafilatura.fetch_url(url)
        text = (
            trafilatura.extract(
                dl,
                include_comments=False,
                include_tables=True, # FONTOS: táblázatok/listák miatt True
                favor_recall=True,
                no_fallback=False,
            )
            if dl
            else None
        )
        paras = []
        if text:
            blocks = [norm_space(b) for b in re.split(r"\n\s*\n", text.replace("\r\n", "\n"))]
            paras = [b for b in blocks if b and not JUNK_RE.search(b)]
        title = ""
        try:
            meta = trafilatura.extract_metadata(dl)
            if meta and getattr(meta, "title", None):
                title = meta.title.strip()
        except Exception:
            pass
        cleaned = clean_and_merge(paras)
        return title, cleaned
    except Exception:
        return "", []

def pick_lead(paras: list[str]) -> str:
    if not paras:
        return ""
    # Keressük az első olyan elemet, ami nem felsorolás
    text = ""
    for p in paras:
        if not p.startswith("•"):
            text = p
            break
    if not text: text = paras[0]
    
    parts = re.split(r"(?<=[.!?…])\s+", text)
    parts = [p.strip() for p in parts if p.strip()]
    if not parts:
        return text
    lead = parts[0]
    if len(parts) >= 2 and len(lead) < 220:
        lead = f"{lead} {parts[1]}"
    return lead.strip()

# ===== AI Kereső Endpoint (Az Apps Script kompatibilitáshoz) =====
@app.post("/chat")
def chat_endpoint(p: ChatPayload):
    if APP_SECRET and (p.secret != APP_SECRET):
        raise HTTPException(status_code=401, detail="Unauthorized")
    
    if not GEMINI_API_KEY:
        return {"sources": []}

    model = genai.GenerativeModel('gemini-1.5-flash')
    try:
        ref_date = datetime.datetime.strptime(p.worksheet, "%Y-%m-%d")
        start_dt = (ref_date - datetime.timedelta(days=7)).strftime("%Y-%m-%d")
        end_dt = (ref_date - datetime.timedelta(days=1)).strftime("%Y-%m-%d")
    except:
        start_dt, end_dt = "last 7 days", "today"

    prompt = f"Magyar gazdasági hírek JSON listaként (title, url) a {p.rovat} témában {start_dt} és {end_dt} között."
    try:
        response = model.generate_content(prompt)
        match = re.search(r'\[.*\]', response.text, re.DOTALL)
        sources = json.loads(match.group()) if match else []
        return {"sources": sources}
    except:
        return {"sources": []}

# ===== Generate Endpoint (Változatlan eredeti logika) =====
@app.post("/generate")
def generate(p: Payload):
    if APP_SECRET and (p.secret != APP_SECRET):
        raise HTTPException(status_code=401, detail="Unauthorized")
    try:
        df = pd.read_csv(csv_url(p.sheet_id, p.worksheet))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Sheet read error: {e}")
    
    df = df.dropna(subset=["Rovat", "Link"])
    df = df[df["Rovat"].astype(str).str.strip() == p.rovat]
    df = df[df["Link"].astype(str).str.startswith(("http://", "https://"), na=False)]
    
    if df.empty:
        raise HTTPException(status_code=404, detail="No data.")

    doc = Document(TEMPLATE_PATH) if os.path.exists(TEMPLATE_PATH) else Document()
    # (Itt az eredeti generáló kódod fut tovább...)
    title_p = doc.add_paragraph()
    try: title_p.style = "Heading 1"
    except: pass
    run = title_p.add_run(f"Weekly News | {p.rovat}")
    run.bold = True
    
    try:
        y, m, d = [int(x) for x in p.worksheet.split("-")]
        monday = datetime.date(y, m, d)
    except: monday = datetime.date.today()
    doc.add_paragraph(hu_date(monday))
    
    add_bm(doc.add_paragraph(), "INTRO")
    rows = df.reset_index(drop=True)
    
    for i, row in rows.iterrows():
        url = str(row["Link"]).strip()
        title, paras = read_paras(url)
        if not title: title = urlparse(url).netloc
        
        intro_line = doc.add_paragraph()
        r = intro_line.add_run(f"{i+1}. {title}")
        r.bold = True
        lead = pick_lead(paras)
        if lead: doc.add_paragraph(lead)
        add_link(doc.add_paragraph(), "read article >>>", f"cikk_{i}")
        doc.add_paragraph("")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    for i, row in rows.iterrows():
        url = str(row["Link"]).strip()
        title, paras = read_paras(url)
        if not title: title = urlparse(url).netloc
        
        ptitle = doc.add_paragraph()
        try: ptitle.style = "Heading 2"
        except: pass
        add_bm(ptitle, f"cikk_{i}")
        ptitle.add_run(title).bold = True
        doc.add_paragraph(f"Source: {urlparse(url).netloc}")
        for para in paras:
            doc.add_paragraph(para)
        add_link(doc.add_paragraph(), "back to intro >>>", "INTRO")
        if i != len(rows) - 1:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"X-Filename": f"BN_{p.rovat}.docx"}
    )
